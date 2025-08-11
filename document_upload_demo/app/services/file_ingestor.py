import mimetypes
from app.models.document import store_document
from PyPDF2 import PdfReader
from docx import Document
import io

def split_text_into_chunks(text, chunk_size=500):
    """Split text into chunks of a specified size."""
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0

    for word in words:
        if current_length + len(word) + 1 > chunk_size:  # +1 for space
            chunks.append(" ".join(current_chunk))
            current_chunk = []
            current_length = 0
        current_chunk.append(word)
        current_length += len(word) + 1

    if current_chunk:
        chunks.append(" ".join(current_chunk))
        
    print(f"Total chunks created: {len(chunks)}")
    return chunks

async def ingest_from_file(data, file):
    content = await file.read()
    mime_type, _ = mimetypes.guess_type(file.filename)
    text = None
    metadata = []

    if mime_type and mime_type.startswith("text"):
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            text = content.decode('latin1', errors='replace')  # Fallback for unusual encodings
    elif mime_type == "application/pdf":
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)

            metadata = []
            chunks = []

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if not page_text:
                    continue

                # Add page metadata
                metadata.append({"page": i+1, "text": page_text})

                # Split this page into chunks
                page_chunks = split_text_into_chunks(page_text)

                # Store chunk with page reference
                for chunk_index, chunk in enumerate(page_chunks):
                    chunks.append({
                        "page": i+1,
                        "chunk_index": chunk_index,
                        "text": chunk
                    })

            chunk_metadata_with_data = {
                **data,
                "file_name": file.filename,
                "mime_type": mime_type,
                "source_metadata": metadata,   # full page texts
                "chunks": chunks               # chunks with page numbers
            }

            await store_document(chunk_metadata_with_data, None)

            return {"status": "success", "chunks_stored": len(chunks)}

        except Exception as e:
            print(f"Error reading PDF: {e}")
            text = "[PDF content could not be extracted]"

    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            doc = Document(io.BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
            metadata = [{"paragraph_index": i, "text": paragraph.text} for i, paragraph in enumerate(doc.paragraphs) if paragraph.text]
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            text = "[DOCX content could not be extracted]"
    else:
        text = f"[Unsupported file type: {file.filename}]"

    print(f"File type: {mime_type}, Content Length: {len(content)}")
    return {"status": "success", "chunks_stored": len(chunks)}